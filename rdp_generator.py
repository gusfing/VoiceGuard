from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_rdp_doc():
    document = Document()

    # Title
    title = document.add_heading('Requirements Definition Package (RDP)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('Project: VoiceGuard – AI Voice Detection API')
    document.add_paragraph('Date: January 23, 2026')
    document.add_paragraph('------------------------------------------------------')

    # 1. Executive Summary
    document.add_heading('1. Executive Summary', level=1)
    document.add_paragraph(
        "Objective: To develop and deploy a REST API capable of analyzing audio samples "
        "in five specified languages (Tamil, English, Hindi, Malayalam, Telugu) and "
        "classifying them as either 'AI_GENERATED' or 'HUMAN'."
    )

    # 2. System Architecture
    document.add_heading('2. System Architecture', level=1)
    document.add_paragraph(
        "The solution utilizes a microservices architecture including an API Gateway, "
        "Preprocessing Unit (Base64 decoding), Inference Engine (Deep Learning Model), "
        "and a Post-Processor for JSON formatting."
    )

    # 3. Functional Requirements
    document.add_heading('3. Functional Requirements', level=1)
    document.add_heading('3.1 Input Specifications', level=2)
    p = document.add_paragraph()
    p.add_run('Protocol: ').bold = True
    p.add_run('HTTPS POST\n')
    p.add_run('Input: ').bold = True
    p.add_run('Base64 encoded MP3 string\n')
    p.add_run('Languages: ').bold = True
    p.add_run('Tamil, English, Hindi, Malayalam, Telugu')

    document.add_heading('3.2 Output Specifications', level=2)
    document.add_paragraph("JSON response containing 'classification' and 'confidence_score'.")

    # 4. API Spec
    document.add_heading('4. API Interface Specification', level=1)
    document.add_paragraph("Endpoint: POST /api/v1/detect")
    
    document.add_heading('Request Body Example:', level=3)
    code_snippet = """{
  "audio_data": "SUQzBAAAAAAAI1RTU0UAAAAPAAADTGF2ZjU..."
}"""
    document.add_paragraph(code_snippet, style='Quote')

    document.add_heading('Response Example:', level=3)
    response_snippet = """{
  "status": "success",
  "data": {
    "classification": "AI_GENERATED",
    "confidence_score": 0.98
  }
}"""
    document.add_paragraph(response_snippet, style='Quote')

    # 5. Tech Stack
    document.add_heading('5. Technical Stack', level=1)
    items = [
        'Language: Python 3.9+',
        'Framework: FastAPI',
        'ML Libraries: PyTorch / TensorFlow',
        'Audio Tools: Librosa, FFMPEG'
    ]
    for item in items:
        document.add_paragraph(item, style='List Bullet')

    # Save
    file_name = 'Voice_Detection_API_RDP.docx'
    document.save(file_name)
    print(f"Document generated successfully: {file_name}")

if __name__ == "__main__":
    create_rdp_doc()
